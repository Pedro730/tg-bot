# ===============  ВАШ БОТ + KEEP-ALIVE  ===============
import os
import logging
import datetime
import hashlib
import asyncio
from pathlib import Path

from docx import Document
from sqlalchemy import create_engine, Column, Integer, String, DateTime
from sqlalchemy.orm import declarative_base, sessionmaker

from telegram import (
    Update, InlineKeyboardButton, InlineKeyboardMarkup, BotCommand
)
from telegram.ext import (
    Application, CommandHandler, MessageHandler, filters,
    ContextTypes, CallbackQueryHandler, ConversationHandler
)
from telegram.helpers import escape

# ---------- KEEP-ALIVE (Flask) ----------
from flask import Flask
import threading

_flask_app = Flask(__name__)

@_flask_app.route("/")
def _ping():
    return "ok", 200

def keep_alive():
    threading.Thread(
        target=lambda: _flask_app.run(host="0.0.0.0", port=8080),
        daemon=True
    ).start()

# ---------- ПЕРЕМЕННЫЕ ОКРУЖЕНИЯ ----------
BOT_TOKEN = os.getenv("BOT_TOKEN")
ADMIN_ID = int(os.getenv("ADMIN_ID", 0))

if not BOT_TOKEN or not ADMIN_ID:
    raise RuntimeError("Укажите BOT_TOKEN и ADMIN_ID в Secrets")

# ---------- БАЗА ДАННЫХ ----------
Path("data").mkdir(exist_ok=True)
DB_URL = f"sqlite:///{Path.cwd() / 'data' / 'users.db'}"
engine = create_engine(DB_URL, echo=False, pool_pre_ping=True)
Base = declarative_base()
SessionLocal = sessionmaker(bind=engine, expire_on_commit=False)

class UserRecord(Base):
    __tablename__ = "users"
    id = Column(Integer, primary_key=True)
    user_id = Column(Integer, unique=True, nullable=False)
    username = Column(String)
    status = Column(String, default="pending")
    requested_at = Column(DateTime, default=datetime.datetime.utcnow)

class SearchHistory(Base):
    __tablename__ = "search_history"
    id = Column(Integer, primary_key=True)
    user_id = Column(Integer, nullable=False)
    username = Column(String)
    query = Column(String, nullable=False)
    timestamp = Column(DateTime, default=datetime.datetime.utcnow)

Base.metadata.create_all(bind=engine)

# ---------- ЛОГИ ----------
logging.basicConfig(
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
    level=logging.INFO
)
logger = logging.getLogger(__name__)

# ---------- DATA.DOCX ----------
DATA_FILE = Path("data.docx")
CHECKSUM_FILE = Path("data.md5")

def create_sample_docx(path: Path):
    doc = Document()
    doc.add_paragraph("Ключевое слово: пример")
    doc.add_paragraph("Описание: Это тестовое описание.")
    doc.save(path)
    logger.info("Создан data.docx")

def _file_checksum(path: Path) -> str:
    if not path.exists():
        return ""
    with open(path, "rb") as f:
        return hashlib.md5(f.read()).hexdigest()

def load_data(file_path: Path) -> dict:
    if not file_path.exists():
        create_sample_docx(file_path)
    try:
        doc = Document(file_path)
    except Exception as e:
        logger.exception("Ошибка чтения docx: %s", e)
        return {}
    data, current_keyword = {}, None
    for para in doc.paragraphs:
        text = para.text.strip()
        if text.startswith("Ключевое слово:"):
            current_keyword = text.replace("Ключевое слово:", "").strip().lower()
        elif text.startswith("Описание:") and current_keyword:
            description = text.replace("Описание:", "").strip()
            data[current_keyword] = description
            current_keyword = None
    return data

def rewrite_data_docx():
    doc = Document()
    for key, desc in DATA.items():
        doc.add_paragraph(f"Ключевое слово: {key}")
        doc.add_paragraph(f"Описание: {desc}")
    doc.save(DATA_FILE)
    CHECKSUM_FILE.write_text(_file_checksum(DATA_FILE))

def _notify_all_approved(app: Application, keys: list[str], action: str):
    """
    Отправляет уведомления всем одобренным пользователям.
    action: 'added', 'edited', 'deleted'
    """
    if action not in {"added", "edited", "deleted"}:
        logger.warning(f"Неверный action: {action}")
        return

    with SessionLocal() as session:
        users = session.query(UserRecord).filter_by(status="approved").all()

    for key in keys:
        if action == "deleted":
            msg = f"🔔 <b>Удалена запись:</b>\n\n<b>{key.capitalize()}</b>"
        else:
            desc = DATA.get(key, "Описание недоступно")
            action_text = "Добавлена новая запись" if action == "added" else "Обновлена запись"
            msg = f"🔔 <b>{action_text}:</b>\n\n<b>{key.capitalize()}</b>\n{desc}"

        for user in users:
            try:
                app.bot.send_message(chat_id=user.user_id, text=msg, parse_mode="HTML")
            except Exception as e:
                logger.warning(f"Не удалось отправить уведомление {user.user_id}: {e}")

def reload_data_and_notify_if_new(app: Application):
    old_keys = set(DATA.keys())
    new_data = load_data(DATA_FILE)
    new_keys = [k for k in new_data.keys() if k not in old_keys]
    DATA.clear()
    DATA.update(new_data)
    if new_keys:
        logger.info(f"Обнаружены новые ключи: {new_keys}")
        _notify_all_approved(app, new_keys, "added")

DATA = load_data(DATA_FILE)
rewrite_data_docx()

# ---------- УТИЛИТЫ ----------
def is_approved(user_id: int) -> bool:
    with SessionLocal() as session:
        user = session.query(UserRecord).filter_by(user_id=user_id).first()
        return bool(user and user.status == "approved")

# ---------- Conversation states ----------
ADD_KEY, ADD_DESC, EDIT_KEY, EDIT_DESC, DELETE_KEY = range(5)
FEEDBACK_TEXT, BROADCAST_TEXT = range(6, 8)

# ---------- Conversation-функции ----------
async def add_start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if update.effective_user.id != ADMIN_ID:
        await update.message.reply_text("❌ Нет доступа.")
        return ConversationHandler.END
    await update.message.reply_text("🔑 Отправь ключевое слово:")
    return ADD_KEY

async def add_key(update: Update, context: ContextTypes.DEFAULT_TYPE):
    context.user_data["add_key"] = update.message.text.strip().lower()
    await update.message.reply_text("📝 Отправь описание:")
    return ADD_DESC

async def add_desc(update: Update, context: ContextTypes.DEFAULT_TYPE):
    key = context.user_data["add_key"]
    desc = update.message.text.strip()
    DATA[key] = desc
    rewrite_data_docx()
    _notify_all_approved(context.application, [key], "added")
    await update.message.reply_text(f"✅ Добавлено:\n<b>{key}</b>\n{desc}", parse_mode="HTML")
    context.user_data.clear()
    return ConversationHandler.END

async def edit_start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if update.effective_user.id != ADMIN_ID:
        await update.message.reply_text("❌ Нет доступа.")
        return ConversationHandler.END
    await update.message.reply_text("🔑 Отправь ключевое слово для редактирования:")
    return EDIT_KEY

async def edit_key(update: Update, context: ContextTypes.DEFAULT_TYPE):
    key = update.message.text.strip().lower()
    if key not in DATA:
        await update.message.reply_text("❌ Запись не найдена.")
        return ConversationHandler.END
    context.user_data["edit_key"] = key
    await update.message.reply_text(f"📝 Текущее описание:\n{DATA[key]}\n\nОтправь новое описание:")
    return EDIT_DESC

async def edit_desc(update: Update, context: ContextTypes.DEFAULT_TYPE):
    key = context.user_data["edit_key"]
    desc = update.message.text.strip()
    DATA[key] = desc
    rewrite_data_docx()
    _notify_all_approved(context.application, [key], "edited")
    await update.message.reply_text(f"✅ Обновлено:\n<b>{key}</b>\n{desc}", parse_mode="HTML")
    context.user_data.clear()
    return ConversationHandler.END

async def del_start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if update.effective_user.id != ADMIN_ID:
        await update.message.reply_text("❌ Нет доступа.")
        return ConversationHandler.END
    await update.message.reply_text("🗑 Отправь ключевое слово для удаления:")
    return DELETE_KEY

async def del_key(update: Update, context: ContextTypes.DEFAULT_TYPE):
    key = update.message.text.strip().lower()
    if key not in DATA:
        await update.message.reply_text("❌ Запись не найдена.")
        return ConversationHandler.END

    deleted_key = key
    del DATA[key]
    rewrite_data_docx()
    _notify_all_approved(context.application, [deleted_key], "deleted")
    await update.message.reply_text(
        f"✅ Запись удалена и уведомления отправлены:\n\n<b>{deleted_key}</b>",
        parse_mode="HTML"
    )
    return ConversationHandler.END

async def cancel(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text("❌ Действие отменено.")
    context.user_data.clear()
    return ConversationHandler.END

# ---------- ОБРАТНАЯ СВЯЗЬ ----------
async def feedback_start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text("📬 Напишите ваше предложение одним сообщением:")
    return FEEDBACK_TEXT

async def feedback_receive(update: Update, context: ContextTypes.DEFAULT_TYPE):
    text = update.message.text
    user = update.effective_user
    await context.bot.send_message(
        chat_id=ADMIN_ID,
        text=f"📩 Feedback от {user.full_name} (@{user.username or '—'} | {user.id}):\n\n{text}"
    )
    await update.message.reply_text("✅ Спасибо, ваше сообщение отправлено администратору!")
    return ConversationHandler.END

# ---------- РАССЫЛКА ----------
async def broadcast_start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text("📢 Отправьте текст для рассылки всем одобренным пользователям:")
    return BROADCAST_TEXT

async def broadcast_send(update: Update, context: ContextTypes.DEFAULT_TYPE):
    text = update.message.text
    with SessionLocal() as session:
        users = session.query(UserRecord).filter_by(status="approved").all()

    sent = 0
    for user in users:
        try:
            await context.bot.send_message(chat_id=user.user_id, text=text, parse_mode="HTML")
            sent += 1
            await asyncio.sleep(0.1)
        except Exception as e:
            logger.warning(f"Не дошло до {user.user_id}: {e}")
    await update.message.reply_text(f"✅ Рассылка завершена. Отправлено: {sent}")
    return ConversationHandler.END

# ---------- ДОБАВЛЕНИЕ ПОЛЬЗОВАТЕЛЕЙ ----------
async def adduser(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if update.effective_user.id != ADMIN_ID:
        return
    if not context.args:
        await update.message.reply_text("❌ Используй: /adduser <user_id>")
        return
    try:
        user_id = int(context.args[0])
    except ValueError:
        await update.message.reply_text("❌ Неверный формат id.")
        return
    with SessionLocal() as session:
        user = session.query(UserRecord).filter_by(user_id=user_id).first()
        if user:
            user.status = "approved"
        else:
            session.add(UserRecord(user_id=user_id, username="N/A", status="approved"))
        session.commit()
    await update.message.reply_text(f"✅ Пользователь {user_id} добавлен и одобрен.")

async def addusers(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if update.effective_user.id != ADMIN_ID:
        return
    if not context.args:
        await update.message.reply_text("❌ Используй: /addusers <id1> <id2> ...")
        return

    added = []
    with SessionLocal() as session:
        for arg in context.args:
            try:
                uid = int(arg.strip())
            except ValueError:
                continue
            user = session.query(UserRecord).filter_by(user_id=uid).first()
            if user:
                user.status = "approved"
            else:
                session.add(UserRecord(user_id=uid, username="N/A", status="approved"))
            added.append(str(uid))
        session.commit()

    await update.message.reply_text(
        f"✅ Добавлено и одобрено: {', '.join(added)}"
    )

# ---------- КОМАНДЫ ----------
async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user = update.effective_user
    with SessionLocal() as session:
        record = session.query(UserRecord).filter_by(user_id=user.id).first()
        if record and record.status == "approved":
            await update.message.reply_text(
                "✅ Добро пожаловать!\n\nОтправьте любое слово для поиска. "
                "В Библиотеке можно найти экстремистов, террористов и других лиц, связанных с экстремизмом и терроризмом, "
                "в том числе запрещённых проповедников и организаций. "
                "Также можно найти запрещённые НС, СЯ и ЯВ. "
                "Поиск осуществляется как по одному-двум словам, так и по его части; "
                "выдаётся 7 наиболее подходящих совпадений, удачи тебе в поисках."
            )
            return
        if record and record.status == "blocked":
            await update.message.reply_text("❌ Вам отказано в доступе.")
            return

        if not record:
            session.add(
                UserRecord(user_id=user.id, username=user.username or "N/A")
            )
            session.commit()

            keyboard = [[InlineKeyboardButton("Одобрить", callback_data=f"approve_{user.id}")]]
            await context.bot.send_message(
                chat_id=ADMIN_ID,
                text=f"📬 Новая заявка:\n"
                     f"ID: {user.id}\n"
                     f"Имя: {user.full_name}\n"
                     f"Username: @{user.username or '—'}",
                reply_markup=InlineKeyboardMarkup(keyboard)
            )
        await update.message.reply_text("📨 Ваша заявка отправлена администратору.")

async def users_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if update.effective_user.id != ADMIN_ID:
        return
    with SessionLocal() as session:
        records = session.query(UserRecord).all()
    if not records:
        await update.message.reply_text("📝 Нет зарегистрированных пользователей.")
        return
    lines, keyboard = [], []
    for r in records:
        status = "✅" if r.status == "approved" else ("❌" if r.status == "blocked" else "⏳")
        lines.append(f"{status} <b>{r.user_id}</b> — {escape(r.username or 'N/A')}")
        keyboard.append([
            InlineKeyboardButton(
                f"{'Заблокать' if r.status == 'approved' else 'Одобрить'} {r.user_id}",
                callback_data=f"toggle_{r.user_id}"
            )
        ])
    await update.message.reply_text(
        "📋 Список пользователей:\n\n" + "\n".join(lines),
        parse_mode="HTML",
        reply_markup=InlineKeyboardMarkup(keyboard)
    )

async def approve_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    if query.from_user.id != ADMIN_ID:
        return
    user_id = int(query.data.split("_")[1])
    with SessionLocal() as session:
        record = session.query(UserRecord).filter_by(user_id=user_id).first()
        if not record:
            await query.edit_message_text("❌ Пользователь не найден")
            return
        record.status = "approved"
        session.commit()
    await query.edit_message_text("✅ Пользователь одобрен")
    try:
        context.bot.send_message(
            chat_id=user_id,
            text="✅ Ваша заявка одобрена! Нажмите /start, чтобы начать."
        )
    except Exception as e:
        logger.warning(f"Не удалось уведомить {user_id}: {e}")

async def toggle_user_status(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    if query.from_user.id != ADMIN_ID:
        return
    user_id = int(query.data.split("_")[1])
    with SessionLocal() as session:
        record = session.query(UserRecord).filter_by(user_id=user_id).first()
        if not record:
            await query.edit_message_text("❌ Пользователь не найден.")
            return
        record.status = "blocked" if record.status == "approved" else "approved"
        session.commit()
    await query.edit_message_text("✅ Статус изменён")

async def history_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if update.effective_user.id != ADMIN_ID:
        return
    with SessionLocal() as session:
        records = session.query(SearchHistory).order_by(SearchHistory.timestamp.desc()).limit(50).all()
    if not records:
        await update.message.reply_text("📭 История поиска пуста.")
        return
    lines = [
        f"{r.timestamp.strftime('%Y-%m-%d %H:%M')} — @{escape(r.username or 'N/A')} — <code>{escape(r.query)}</code>"
        for r in records
    ]
    await update.message.reply_text(
        "📋 История поиска (последние 50):\n\n" + "\n".join(lines),
        parse_mode="HTML"
    )

async def stats_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if update.effective_user.id != ADMIN_ID:
        return
    with SessionLocal() as session:
        total_users = session.query(UserRecord).count()
        approved_users = session.query(UserRecord).filter_by(status="approved").count()
        total_searches = session.query(SearchHistory).count()
        today_searches = session.query(SearchHistory).filter(
            SearchHistory.timestamp >= datetime.datetime.utcnow().date()
        ).count()
    await update.message.reply_text(
        f"📊 Статистика:\n\n"
        f"👥 Всего пользователей: {total_users}\n"
        f"✅ Одобрено: {approved_users}\n"
        f"🔍 Всего поисков: {total_searches}\n"
        f"📅 За сегодня: {today_searches}"
    )

async def list_entries(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if update.effective_user.id != ADMIN_ID:
        return
    if not DATA:
        await update.message.reply_text("📭 База пуста.")
        return
    keyboard = []
    for key in sorted(DATA):
        keyboard.append([
            InlineKeyboardButton(f"✏️ {key}", callback_data=f"e_{key}"),
            InlineKeyboardButton(f"🗑️ {key}", callback_data=f"d_{key}")
        ])
    await update.message.reply_text(
        "📋 Выберите запись:",
        reply_markup=InlineKeyboardMarkup(keyboard)
    )

async def list_button(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    if query.from_user.id != ADMIN_ID:
        return
    cmd, key = query.data.split("_", 1)
    if cmd == "e":
        context.user_data["edit_key"] = key
        await query.edit_message_text(
            f"📝 Текущее описание:\n{DATA[key]}\n\nОтправь новое описание:"
        )
        return EDIT_DESC
    elif cmd == "d":
        deleted_key = key
        del DATA[key]
        rewrite_data_docx()
        _notify_all_approved(context.application, [deleted_key], "deleted")
        await query.edit_message_text(
            f"✅ Запись удалена и уведомления отправлены:\n\n<b>{deleted_key}</b>",
            parse_mode="HTML"
        )

async def handle_message(update: Update, context: ContextTypes.DEFAULT_TYPE):
    try:
        if not is_approved(update.effective_user.id):
            await update.message.reply_text("❌ У вас нет доступа.")
            return
        query = update.message.text.strip().lower()
        if not query:
            await update.message.reply_text("🔍 Пустой запрос.")
            return
        with SessionLocal() as session:
            session.add(SearchHistory(
                user_id=update.effective_user.id,
                username=update.effective_user.username or "N/A",
                query=query
            ))
            session.commit()

        matches = [(k, v) for k, v in DATA.items() if query in k or k in query][:7]
        if not matches:
            await update.message.reply_text("🔍 Ничего не найдено.")
            return
        lines = [f"<b>{escape(k.capitalize())}</b>\n{escape(v)}" for k, v in matches]
        text = "\n\n".join(lines)
        if len(matches) == 7:
            text += "\n\n<i>Показано первые 7 совпадений</i>"
        await update.message.reply_text(text, parse_mode="HTML")
    except Exception as e:
        logger.exception("Ошибка в handle_message: %s", e)
        await update.message.reply_text("⚠️ Произошла ошибка, попробуйте позже.")

async def unknown(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text("🤷‍♂️ Ну ну ну, разогнался... Нажми /start")

# ---------- HANDLERS ----------
conv_add = ConversationHandler(
    entry_points=[CommandHandler("add", add_start, filters=filters.User(user_id=ADMIN_ID))],
    states={
        ADD_KEY: [MessageHandler(filters.TEXT & ~filters.COMMAND, add_key)],
        ADD_DESC: [MessageHandler(filters.TEXT & ~filters.COMMAND, add_desc)]
    },
    fallbacks=[CommandHandler("cancel", cancel)]
)

conv_edit = ConversationHandler(
    entry_points=[CommandHandler("edit", edit_start, filters=filters.User(user_id=ADMIN_ID))],
    states={
        EDIT_KEY: [MessageHandler(filters.TEXT & ~filters.COMMAND, edit_key)],
        EDIT_DESC: [MessageHandler(filters.TEXT & ~filters.COMMAND, edit_desc)]
    },
    fallbacks=[CommandHandler("cancel", cancel)]
)

conv_del = ConversationHandler(
    entry_points=[CommandHandler("del", del_start, filters=filters.User(user_id=ADMIN_ID))],
    states={
        DELETE_KEY: [MessageHandler(filters.TEXT & ~filters.COMMAND, del_key)]
    },
    fallbacks=[CommandHandler("cancel", cancel)]
)

conv_feedback = ConversationHandler(
    entry_points=[CommandHandler("feedback", feedback_start)],
    states={
        FEEDBACK_TEXT: [MessageHandler(filters.TEXT & ~filters.COMMAND, feedback_receive)]
    },
    fallbacks=[CommandHandler("cancel", cancel)]
)

conv_broadcast = ConversationHandler(
    entry_points=[CommandHandler("broadcast", broadcast_start, filters=filters.User(user_id=ADMIN_ID))],
    states={
        BROADCAST_TEXT: [MessageHandler(filters.TEXT & ~filters.COMMAND, broadcast_send)]
    },
    fallbacks=[CommandHandler("cancel", cancel)]
)

async def post_init(app: Application):
    commands = [
        BotCommand("start", "Начать работу"),
        BotCommand("feedback", "Отправить предложение админу")
    ]
    if ADMIN_ID:
        commands.extend([
            BotCommand("adduser", "Добавить/восстановить пользователя по id"),
            BotCommand("addusers", "Добавить и одобрить нескольких по id"),
            BotCommand("add", "Добавить запись"),
            BotCommand("edit", "Изменить запись"),
            BotCommand("del", "Удалить запись"),
            BotCommand("list", "Список записей"),
            BotCommand("history", "История поиска"),
            BotCommand("stats", "Статистика"),
            BotCommand("users", "Список пользователей"),
            BotCommand("broadcast", "Рассылка всем (админ)"),
            BotCommand("cancel", "Отменить")
        ])
    await app.bot.set_my_commands(commands)

def main():
    keep_alive()
    application = Application.builder().token(BOT_TOKEN).post_init(post_init).build()

    # общедоступные
    application.add_handler(CommandHandler("start", start))
    application.add_handler(conv_feedback)

    # админские
    application.add_handler(CommandHandler("adduser", adduser, filters=filters.User(user_id=ADMIN_ID)))
    application.add_handler(CommandHandler("addusers", addusers, filters=filters.User(user_id=ADMIN_ID)))
    application.add_handler(CommandHandler("users", users_command, filters=filters.User(user_id=ADMIN_ID)))
    application.add_handler(CommandHandler("history", history_command, filters=filters.User(user_id=ADMIN_ID)))
    application.add_handler(CommandHandler("stats", stats_command, filters=filters.User(user_id=ADMIN_ID)))
    application.add_handler(CommandHandler("list", list_entries, filters=filters.User(user_id=ADMIN_ID)))
    application.add_handler(conv_add)
    application.add_handler(conv_edit)
    application.add_handler(conv_del)
    application.add_handler(conv_broadcast)

    # callback-и
    application.add_handler(CallbackQueryHandler(approve_callback, pattern="^approve_"))
    application.add_handler(CallbackQueryHandler(toggle_user_status, pattern="^toggle_"))
    application.add_handler(CallbackQueryHandler(list_button, pattern="^[ed]_"))

    # текстовые сообщения
    application.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle_message))
    application.add_handler(MessageHandler(filters.COMMAND, unknown))

    logger.info("✅ Бот запущен")
    application.run_polling()

if __name__ == "__main__":
    main()
